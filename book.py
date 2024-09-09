from docx import Document
from dotenv import load_dotenv
import logging
from openai import OpenAI
import os
from pathlib import Path
from pydantic import BaseModel
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import time


load_dotenv(Path("./api_key.env"))
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Chapter(BaseModel):
    title: str
    description: str

class Chapters(BaseModel):
    chapters: list[Chapter]

class Subsection(BaseModel):
    title: str
    description: str

class Subsections(BaseModel):
    subsections: list[Subsection]

class SubsectionContent(BaseModel):
    content: str

class BookOpenAI:
    def __init__(self, model_name="gpt-4o-mini"):
        self.model_name = model_name
        self.client = OpenAI()
        self.messages = []
        self.chapters = {}
        self.title = ""
        self.description = ""
        self.writing_style = ""

    def generate_chapters(self, title, description,writing_style):
        logging.info("Starting chapter generation...")
        self.title = title
        self.description = description
        self.writing_style = writing_style

        self.messages = []  # Clear previous messages

        start_time = time.time()
        completion = self.client.beta.chat.completions.parse(
            model=self.model_name,
            messages=[
                {"role": "system", "content": "Generate the book chapters using the book title and description."},
                {"role": "user", "content": f"Book Title: '{title}'\nDescription: '{description}'"},
            ],
            response_format=Chapters,
        )
        end_time = time.time()
        logging.info(f"Chapters generated in {end_time - start_time:.2f} seconds")

        chapters = completion.choices[0].message.parsed.chapters
        logging.debug(f"Chapters: {chapters}")

        # Store chapters and their descriptions
        for chapter in chapters:
            self.chapters[chapter.title] = {"description": chapter.description, "subsections": {}}
        self.messages.append({"role": "assistant", "content": " ".join([chapter.title for chapter in chapters])})

        return chapters

    def generate_subsections(self, chapters):
        for chapter in chapters:
            logging.info(f"Generating subsections for chapter: {chapter}")
            start_time = time.time()
            completion = self.client.beta.chat.completions.parse(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Generate the chapter subsections."},
                    {"role": "user", "content": f"Chapter Title: {chapter}\nChapter Description: {self.chapters[chapter]['description']}: generate subsections."},
                ],
                response_format=Subsections,
            )
            end_time = time.time()
            logging.info(f"Subsections for chapter '{chapter}' generated in {end_time - start_time:.2f} seconds")

            subsections = completion.choices[0].message.parsed.subsections
            self.messages.extend(subsections)

            for subsection in subsections:
                self.chapters[chapter]["subsections"][subsection.title] = {"description": subsection.description}
            
            logging.info(f"Subsections generated for chapter: {chapter}")

        logging.info("All subsections have been generated.")
        return chapters

    def generate_content(self):
        logging.info("Starting content generation...")
        start_time = time.time()
        for chapter_title in self.chapters.keys():
            chapter_description = self.chapters[chapter_title]["description"]
            logging.info("Starting {chapter_title} generation...")
            for subsection_title in self.chapters[chapter_title]["subsections"].keys():
                subsection_description = self.chapters[chapter_title]["subsections"][subsection_title]["description"]

                subsection_prompt = f"""\
                                    \nBook Title: '{self.title}'
                                    \nBook Description: '{self.description}'
                                    \nActual Chapter Title: '{chapter_title}'   
                                    \nActual Chapter Description: '{chapter_description}'
                                    \nActual Subsection Title: '{subsection_title}'
                                    \nActual Subsection Description: '{subsection_description}'\
                                    \nWriting style: '{self.writing_style}'\
                                    """

                completion = self.client.beta.chat.completions.parse(
                    model=self.model_name,
                    messages=[
                        {"role": "system", "content": "Given the information, generate the Actual Subsection content."},
                        {"role": "user", "content": subsection_prompt},
                    ],
                    response_format=SubsectionContent,
                )

                subsection_content = completion.choices[0].message.parsed.content
                self.chapters[chapter_title]["subsections"][subsection_title]["content"] = subsection_content
                self.messages.append({"role": "assistant", "content": subsection_content})

        end_time = time.time()
        logging.info(f"Content generation completed in {end_time - start_time:.2f} seconds")

    def save_as_pdf(self, filename):
        logging.info(f"Saving book as PDF: {filename}")
        start_time = time.time()
        c = canvas.Canvas(filename, pagesize=letter)
        width, height = letter
        y_position = height - 50
        for chapter_title, chapter_data in self.chapters.items():
            c.drawString(35, y_position, f"{chapter_title}")
            y_position -= 20
            
            for subsection in chapter_data["subsections"].values():
                c.drawString(50, y_position, f"{subsection['content']}")
                y_position -= 20
                
                if y_position < 40:
                    c.showPage()
                    y_position = height - 50

        c.save()
        end_time = time.time()
        logging.info(f"PDF saved in {end_time - start_time:.2f} seconds")

    def save_as_txt(self, filename):
        logging.info(f"Saving book as TXT: {filename}")
        start_time = time.time()
        with open(filename, 'w') as f:
            for chapter_title, chapter_data in self.chapters.items():
                f.write(f"{chapter_title}\n")
                for subsection in chapter_data['subsections'].values():
                    f.write(f"{subsection['content']}\n\n")
        end_time = time.time()
        logging.info(f"TXT file saved in {end_time - start_time:.2f} seconds")

    def save_as_docx(self, filename):
        logging.info(f"Saving book as DOCX: {filename}")
        start_time = time.time()
        doc = Document()
        for chapter_title, chapter_data in self.chapters.items():
            doc.add_heading(chapter_title, level=1)

            for subsection in chapter_data["subsections"].values():
                doc.add_paragraph(subsection['content'])

        doc.save(filename)
        end_time = time.time()
        logging.info(f"DOCX file saved in {end_time - start_time:.2f} seconds")
