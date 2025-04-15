# book_openai.py
# Version with fixes for TOC literal_eval and AttributeError

from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
import logging
from openai import OpenAI
import os
from pathlib import Path
from pydantic import BaseModel
from reportlab.lib.pagesizes import letter
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
import time
import re
import traceback

# --- Load environment variables and configure logging ---
load_dotenv(dotenv_path=Path("./.env"), verbose=True)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    logging.error("ERROR: OPENAI_API_KEY not found in environment after loading .env file.")
else:
    logging.info("OpenAI API Key loaded successfully (first few chars): %s", api_key[:4] + "..." if api_key else "None")

# --- Pydantic Models ---
class Language(BaseModel): language: str
class Translation(BaseModel): translation: str
class Chapter(BaseModel): title: str; description: str
class Chapters(BaseModel): chapters: list[Chapter]
class Subsection(BaseModel): title: str; description: str
class Subsections(BaseModel): subsections: list[Subsection]
class SubsectionContent(BaseModel): content: str

# --- Helper Functions ---
def clean_content(content):
    if not isinstance(content, str): return ""
    content = re.sub(r'^###.*$', '', content, flags=re.MULTILINE)
    content = re.sub(r'\n\s*\n', '\n', content)
    content = content.strip()
    return content

def strip_chapter_prefix(chapter_title):
    if not isinstance(chapter_title, str): return ""
    return re.sub(r'^Chapter\s*\d+:\s*', '', chapter_title, flags=re.IGNORECASE).strip()

# --- PDF Generation Helper Functions ---
def add_page_number(canvas_obj, doc_obj):
    """Add page number to the footer of each page."""
    canvas_obj.saveState()
    page_number_text = f"{doc_obj.page}"
    canvas_obj.setFont('Helvetica', 10)
    page_width = letter[0]
    canvas_obj.drawCentredString(page_width / 2.0, 0.5 * inch, page_number_text)
    canvas_obj.restoreState()

# --- CORRECTED MyDocTemplate Class ---
class MyDocTemplate(BaseDocTemplate):
    """Custom Document Template with TOC and Page Numbers."""
    def __init__(self, filename, **kw):
        """Initialize the document template, frame, page template, and TOC object."""
        # Call parent __init__
        super().__init__(filename, **kw)

        # Define the main frame with 1-inch margins
        main_frame = Frame(
            x1=1 * inch, y1=1 * inch,
            width=letter[0] - 2 * inch, height=letter[1] - 2 * inch,
            id='main_frame',
            leftPadding=0, bottomPadding=0, # Explicitly set padding if needed
            rightPadding=0, topPadding=0
        )

        # Create a page template using the frame and add page numbering
        main_template = PageTemplate(id='main', frames=[main_frame], onPage=add_page_number)
        self.addPageTemplates([main_template]) # Add the template

        # --- FIX: Initialize the TableOfContents object ---
        self.toc = TableOfContents()
        # Configure TOC appearance
        self.toc.levelStyles = [
            ParagraphStyle(
                name='TOCHeading1', fontName='Helvetica-Bold', fontSize=14,
                leftIndent=20, firstLineIndent=-20, spaceBefore=6, leading=16
            ),
            ParagraphStyle(
                name='TOCHeading2', fontName='Helvetica', fontSize=12,
                leftIndent=40, firstLineIndent=-20, spaceBefore=4, leading=14
            ),
        ]
        # --- End FIX ---

    def afterFlowable(self, flowable):
        """Registers TOC entries and handles bookmarking."""
        if isinstance(flowable, Paragraph):
            text = flowable.getPlainText()
            style = flowable.style.name
            # Generate a unique key for bookmarking (simple approach)
            bookmark_key = f"{style}_{text[:20]}".replace(" ","_") # Basic unique key

            if style == 'ChapterTitle':
                # Format text for TOC display
                match = re.match(r'Chapter\s*(\d+):', text, re.IGNORECASE)
                level_text = match.group(1) + ". " + strip_chapter_prefix(text) if match else text
                # Notify TOC mechanism (Level 0 for chapters) - CORRECTED (no 4th element)
                self.notify('TOCEntry', (0, level_text, self.page))
                # Add PDF outline entry and bookmark destination
                self.canv.bookmarkPage(bookmark_key)
                self.canv.addOutlineEntry(level_text, bookmark_key, level=0, closed=0)

            elif style == 'SubsectionTitle':
                # Notify TOC mechanism (Level 1 for subsections) - CORRECTED (no 4th element)
                self.notify('TOCEntry', (1, text, self.page))
                # Add PDF outline entry and bookmark destination
                self.canv.bookmarkPage(bookmark_key)
                self.canv.addOutlineEntry(text, bookmark_key, level=1, closed=0)


# --- Main Book Generation Class BookOpenAI ---
class BookOpenAI:
    def __init__(self, model_name="gpt-4.1-nano"):
        """Initialize the BookOpenAI instance."""
        self.model_name = model_name
        try:
            self.client = OpenAI() # Assumes OPENAI_API_KEY is set in environment
            # Simple check if client was created (optional)
            # self.client.models.list(limit=1)
            logging.info("OpenAI client initialized successfully.")
        except Exception as e:
            logging.error(f"Failed to initialize OpenAI client: {e}", exc_info=True)
            # Depending on requirements, you might want to raise the error
            # raise ConnectionError(f"Failed to initialize OpenAI client: {e}") from e
            self.client = None # Ensure client is None if init fails

        self.chapters = {} # Main data structure
        self.title = ""
        self.description = ""
        self.writing_style = ""
        self.target_language = "en"

    # --- Language and Chapter/Subsection Generation Logic ---
    def extract_language(self, text):
        """Extract the primary language from the text using OpenAI."""
        if not self.client: return "en" # Return default if client failed
        if not text: return "en" # Default for empty text
        prompt = f"Identify the primary language of the following text and return only its two-letter ISO 639-1 code (e.g., 'en', 'es', 'fr', 'de'). Text: '{text}'"
        messages = [{"role": "system", "content": "You are a language ID assistant. Respond with only the two-letter ISO 639-1 code."}, {"role": "user", "content": prompt}]
        try:
            completion = self.client.chat.completions.create(model=self.model_name, messages=messages, temperature=0.1, max_tokens=10)
            language_code = completion.choices[0].message.content.strip().lower()
            if re.match(r'^[a-z]{2}$', language_code):
                 logging.info(f"Extracted language code: {language_code}")
                 return language_code
            else:
                 logging.warning(f"Unexpected language format: '{language_code}'. Defaulting to 'en'.")
                 return "en"
        except Exception as e:
            logging.error(f"Language extraction failed: {e}", exc_info=True); return "en" # Default on error

    def generate_chapters(self, title, description, writing_style):
        """Generate chapters, detect language, set internal state."""
        if not self.client: logging.error("OpenAI client not available."); return None
        logging.info("Starting chapter generation...")
        self.title = title; self.description = description; self.writing_style = writing_style; self.chapters = {}
        self.target_language = self.extract_language(description or title) # Use description or title
        logging.info(f"Using target language: {self.target_language}")

        start_time = time.time()
        system_message = f"Generate a comprehensive list of chapter titles and brief descriptions for a book titled '{title}' about '{description}' in {self.target_language}. Style: {writing_style}. Respond strictly in the required Pydantic format."
        user_prompt = f"Book Title: '{self.title}'\nDescription: '{self.description}'\nStyle: '{self.writing_style}'\nGenerate chapters."
        try:
            completion = self.client.beta.chat.completions.parse(model=self.model_name, messages=[{"role": "system", "content": system_message}, {"role": "user", "content": user_prompt}], response_format=Chapters, max_tokens=2000)
            generated_chapters_list = completion.choices[0].message.parsed.chapters
            if not generated_chapters_list: logging.error("Chapter gen resulted in empty list."); return None

            logging.info(f"Chapters generated ({len(generated_chapters_list)}) in {time.time() - start_time:.2f}s")
            self.chapters = {}
            for chapter_obj in generated_chapters_list:
                cleaned_title = strip_chapter_prefix(chapter_obj.title) or f"Untitled Chapter {len(self.chapters) + 1}"
                self.chapters[cleaned_title] = {"description": chapter_obj.description, "subsections": {}}
            logging.debug(f"Stored chapters: {list(self.chapters.keys())}")
            return generated_chapters_list
        except Exception as e: logging.error(f"Failed to generate/parse chapters: {e}", exc_info=True); return None

    def generate_subsections(self, chapters_list, progress_callback=None):
        """Generate subsections, invoking callback."""
        if not self.client: logging.error("OpenAI client not available."); return
        if not chapters_list: logging.warning("No chapters provided."); return
        logging.info("Generating subsections for all chapters...")
        total_start_time = time.time(); total_chapters = len(chapters_list)

        for i, chapter_obj in enumerate(chapters_list):
            chapter_title_key = strip_chapter_prefix(chapter_obj.title) or f"Untitled Chapter {i+1}"
            if chapter_title_key not in self.chapters:
                logging.warning(f"Chapter key '{chapter_title_key}' (from obj '{chapter_obj.title}') not found. Skipping.")
                continue

            if progress_callback:
                try: progress_callback(i, total_chapters)
                except Exception as cb_err: logging.error(f"Err in subsection progress cb: {cb_err}", exc_info=True)

            logging.info(f"Generating subsections for Ch {i+1}/{total_chapters}: '{chapter_title_key}'")
            start_time = time.time()
            chapter_data = self.chapters[chapter_title_key]
            system_message = f"Generate logical subsection titles & descriptions for chapter '{chapter_title_key}' ({chapter_data['description']}) of book '{self.title}'. Language: {self.target_language}. Pydantic format."
            user_prompt = f"Chapter: '{chapter_title_key}'\nDescription: '{chapter_data['description']}'\nGenerate subsections."
            try:
                completion = self.client.beta.chat.completions.parse(model=self.model_name, messages=[{"role": "system", "content": system_message}, {"role": "user", "content": user_prompt}], response_format=Subsections, max_tokens=1500)
                subsections_list = completion.choices[0].message.parsed.subsections
                self.chapters[chapter_title_key]["subsections"] = {}
                if not subsections_list:
                    logging.warning(f"No subsections generated for '{chapter_title_key}'.")
                else:
                    for sub_obj in subsections_list:
                        sub_title = sub_obj.title.strip() or f"Untitled Subsection {len(self.chapters[chapter_title_key]['subsections']) + 1}"
                        self.chapters[chapter_title_key]["subsections"][sub_title] = {"description": sub_obj.description, "content": None}
                    logging.info(f"Subsections for '{chapter_title_key}' ({len(subsections_list)}) generated in {time.time() - start_time:.2f}s")
            except Exception as e:
                logging.error(f"Failed gen/parse subsections for '{chapter_title_key}': {e}", exc_info=True)
                self.chapters[chapter_title_key]["subsections"] = {}

        logging.info(f"All subsection generation finished in {time.time() - total_start_time:.2f} seconds.")

    def generate_content(self, progress_callback=None):
        """Generate content, invoking callback."""
        if not self.client: logging.error("OpenAI client not available."); return
        logging.info("Starting content generation..."); overall_start_time = time.time()
        if not self.chapters: logging.error("Cannot generate content: No chapters."); return

        total_chapters = len(self.chapters)
        total_subsections = sum(len(data.get("subsections", {})) for data in self.chapters.values())
        processed_subsections = 0
        logging.info(f"Total chapters: {total_chapters}, Total subsections: {total_subsections}")

        if total_subsections == 0:
            logging.warning("No subsections found. Skipping content generation.")
            if progress_callback: 
                try: 
                    progress_callback(0, 0, 0, total_chapters, 0, 0)
                except Exception as cb_err:
                    logging.error(f"Err in empty content cb: {cb_err}", exc_info=True)
            return

        for i, (chapter_title_key, chapter_data) in enumerate(self.chapters.items()):
            subsections_dict = chapter_data.get("subsections", {})
            num_subsections_in_chapter = len(subsections_dict)
            logging.info(f"--- Generating content for Ch {i+1}/{total_chapters}: '{chapter_title_key}' ({num_subsections_in_chapter} subs) ---")
            if not subsections_dict: continue

            for j, (subsection_title_key, subsection_data) in enumerate(subsections_dict.items()):
                processed_subsections += 1
                if progress_callback:
                    try: progress_callback(processed_subsections, total_subsections, i, total_chapters, j, num_subsections_in_chapter)
                    except Exception as cb_err: logging.error(f"Err in content progress cb: {cb_err}", exc_info=True)

                logging.info(f"Generating Sub {j+1}/{num_subsections_in_chapter}: '{subsection_title_key}' (Overall {processed_subsections}/{total_subsections})")
                start_time = time.time()
                system_message = f"You are writing subsection '{subsection_title_key}' for chapter '{chapter_title_key}' of book '{self.title}'. Language: {self.target_language}. Style: '{self.writing_style}'. Use Markdown **bold**. Generate detailed content for *this subsection only*. Respond strictly with the content in Pydantic format."
                user_prompt = f"Context:\nBook: '{self.title}'\nChapter {i+1}: '{chapter_title_key}' ({chapter_data.get('description', 'N/A')})\nSubsection: '{subsection_title_key}' ({subsection_data.get('description', 'N/A')})\nGenerate content:"
                try:
                    completion = self.client.beta.chat.completions.parse(model=self.model_name, messages=[{"role": "system", "content": system_message}, {"role": "user", "content": user_prompt}], response_format=SubsectionContent, temperature=0.6, max_tokens=4000)
                    generated_content = completion.choices[0].message.parsed.content
                    self.chapters[chapter_title_key]["subsections"][subsection_title_key]["content"] = generated_content
                    logging.info(f"Content for '{subsection_title_key}' gen in {time.time() - start_time:.2f}s.")
                except Exception as e:
                    logging.error(f"Failed gen content for '{subsection_title_key}': {e}", exc_info=True)
                    self.chapters[chapter_title_key]["subsections"][subsection_title_key]["content"] = f"Error: Content generation failed. {e}"

        logging.info(f"Content gen for {processed_subsections} subs completed in {time.time() - overall_start_time:.2f}s.")

    # --- Saving Methods ---
    def save_as_txt(self, filename):
        """Save the generated book as a plain text file (.txt)."""
        logging.info(f"Saving book as TXT: {filename}")
        if not self.chapters: logging.error("Cannot save TXT: No chapters."); raise ValueError("No chapters generated.")
        full_content = f"Book Title: {self.title}\n{'=' * (len(self.title) + 12)}\n\n"
        for i, (chapter_title_key, chapter_data) in enumerate(self.chapters.items()):
            full_content += f"--- Chapter {i+1}: {chapter_title_key} ---\n\n"
            subsections = chapter_data.get("subsections", {})
            if not subsections: full_content += "(No subsections generated)\n\n"; continue
            for sub_title_key, sub_data in subsections.items():
                full_content += f"--- Subsection: {sub_title_key} ---\n"
                content = sub_data.get('content', 'Content not generated.')
                cleaned_content = clean_content(content)
                full_content += f"{cleaned_content}\n\n"
            full_content += "\n"
        try:
            with open(filename, 'w', encoding='utf-8') as f: f.write(full_content)
            logging.info("TXT file saved successfully.")
        except IOError as e: logging.error(f"Error saving TXT file '{filename}': {e}", exc_info=True); raise
        except Exception as e: logging.error(f"Unexpected error during TXT save: {e}", exc_info=True); raise

    def save_as_docx(self, filename):
        """Save the generated book as a Microsoft Word document (.docx)."""
        logging.info(f"Saving book as DOCX: {filename}")
        if not self.chapters: logging.error("Cannot save DOCX: No chapters."); raise ValueError("No chapters generated.")
        document = Document(); document.add_heading(self.title, level=0)
        for i, (chapter_title_key, chapter_data) in enumerate(self.chapters.items()):
            document.add_heading(f"Chapter {i+1}: {chapter_title_key}", level=1)
            subsections = chapter_data.get("subsections", {})
            if not subsections:
                document.add_paragraph("(No subsections generated)")
                if i < len(self.chapters) - 1: document.add_page_break()
                continue
            for sub_title_key, sub_data in subsections.items():
                document.add_heading(sub_title_key, level=2)
                content = sub_data.get('content', 'Content not generated.')
                cleaned_content = clean_content(content)
                paragraphs = cleaned_content.split('\n')
                for para_text in paragraphs:
                    if para_text.strip(): document.add_paragraph(para_text.strip())
                document.add_paragraph() # Spacing
            if i < len(self.chapters) - 1: document.add_page_break()
        try:
            document.save(filename)
            logging.info("DOCX file saved successfully.")
        except IOError as e: logging.error(f"Error saving DOCX file '{filename}': {e}", exc_info=True); raise
        except Exception as e: logging.error(f"Unexpected error during DOCX save: {e}", exc_info=True); raise

    def save_as_pdf(self, filename):
        """Save the generated book as PDF with TOC and basic formatting."""
        logging.info(f"Saving book as PDF: {filename}")
        if not self.chapters: logging.error("Cannot save PDF: No chapters."); raise ValueError("No chapters generated.")
        start_time = time.time()
        # Instantiate the corrected MyDocTemplate
        doc = MyDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()

        # Define PDF Styles
        styles.add(ParagraphStyle(name='TitleCentered', parent=styles['Title'], alignment=TA_CENTER, spaceAfter=24))
        styles.add(ParagraphStyle(name='TOCHeader', parent=styles['h1'], alignment=TA_LEFT, spaceAfter=12, fontSize=16))
        styles.add(ParagraphStyle(name='ChapterTitle', parent=styles['h1'], fontSize=18, leading=22, spaceBefore=12, spaceAfter=12, alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='SubsectionTitle', parent=styles['h2'], fontSize=14, leading=18, spaceBefore=10, spaceAfter=6, alignment=TA_LEFT))
        styles.add(ParagraphStyle(name='Content', parent=styles['BodyText'], fontSize=12, leading=15, spaceAfter=12, alignment=TA_JUSTIFY))

        story = []
        # Title page element
        story.append(Paragraph(self.title, styles['TitleCentered']))
        story.append(Spacer(1, 0.5*inch))

        # TOC Header and placeholder (doc.toc is now guaranteed to exist)
        story.append(Paragraph("Table of Contents", styles['TOCHeader']))
        story.append(doc.toc) # Add the TOC object to the story
        story.append(PageBreak())

        # Add Chapters and Subsections to Story
        first_chapter = True
        for i, (chapter_title_key, chapter_data) in enumerate(self.chapters.items()):
            if not first_chapter: story.append(PageBreak())
            else: first_chapter = False

            # Chapter Title - NOTE: This title IS passed to afterFlowable for TOC creation
            chapter_para = Paragraph(f"Chapter {i+1}: {chapter_title_key}", styles['ChapterTitle'])
            story.append(chapter_para)

            subsections = chapter_data.get("subsections", {})
            if not subsections:
                story.append(Paragraph("(No subsections generated)", styles['Content']))
                continue

            # Process Subsections
            for subsection_title_key, subsection_data in subsections.items():
                 # Subsection Title - NOTE: This title IS passed to afterFlowable for TOC creation
                subsection_para = Paragraph(subsection_title_key, styles['SubsectionTitle'])
                story.append(subsection_para)

                # Content processing
                raw_content = subsection_data.get('content', 'Content not generated.')
                cleaned_content = clean_content(raw_content)
                # Convert basic **markdown bold** to <b>reportlab bold</b>
                formatted_content = re.sub(r'\*\*(.*?)\*\*', r'<b>\\1</b>', cleaned_content, flags=re.DOTALL)
                 # Replace newlines with <br/> tags for PDF paragraphs
                content_for_pdf = formatted_content.replace('\n', '<br/>')
                # Create content Paragraph
                content_para = Paragraph(content_for_pdf, styles['Content'])
                story.append(content_para)

        # Build the PDF document
        try:
            doc.multiBuild(story) # This triggers the afterFlowable calls
            logging.info(f"PDF saved successfully in {time.time() - start_time:.2f} seconds.")
        except Exception as e:
            logging.error(f"Error during PDF build for '{filename}': {e}", exc_info=True)
            raise # Re-raise the exception to be caught by Gradio
